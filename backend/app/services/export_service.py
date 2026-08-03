"""
Serviço de exportação de petições/documentos para Word (.docx) e PDF.

Usa python-docx para Word e reportlab para PDF, ambos amplamente
suportados e sem dependências externas adicionais.
"""

import io
import logging
import re
from typing import Optional

logger = logging.getLogger(__name__)


def export_docx(titulo: str, conteudo: str) -> bytes:
    """Gera um arquivo .docx e retorna os bytes."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading(titulo, level=1)

    for bloco in conteudo.split("\n\n"):
        bloco = bloco.strip()
        if not bloco:
            continue
        # Linhas que parecem títulos Markdown.
        if bloco.startswith("# "):
            doc.add_heading(bloco[2:].strip(), level=1)
        elif bloco.startswith("## "):
            doc.add_heading(bloco[3:].strip(), level=2)
        elif bloco.startswith("### "):
            doc.add_heading(bloco[4:].strip(), level=3)
        else:
            doc.add_paragraph(bloco)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def export_pdf(titulo: str, conteudo: str) -> bytes:
    """Gera um arquivo .pdf e retorna os bytes."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=2.5 * cm, rightMargin=2.5 * cm, topMargin=2.5 * cm, bottomMargin=2.5 * cm)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleLEX", parent=styles["Title"], fontSize=16, leading=20, spaceAfter=18, textColor="#0B1F3A"
    )
    heading_style = ParagraphStyle(
        "HeadingLEX", parent=styles["Heading2"], fontSize=12, leading=16, spaceBefore=12, spaceAfter=6, textColor="#102B52"
    )
    body_style = ParagraphStyle(
        "BodyLEX", parent=styles["BodyText"], fontSize=10.5, leading=15, spaceAfter=8
    )

    story = [Paragraph(_escape_html(titulo), title_style), Spacer(1, 0.5 * cm)]

    for bloco in conteudo.split("\n\n"):
        bloco = bloco.strip()
        if not bloco:
            continue
        if bloco.startswith("# "):
            story.append(Paragraph(_escape_html(bloco[2:].strip()), heading_style))
        elif bloco.startswith("## "):
            story.append(Paragraph(_escape_html(bloco[3:].strip()), heading_style))
        elif bloco.startswith("### "):
            story.append(Paragraph(_escape_html(bloco[4:].strip()), heading_style))
        else:
            texto = _bold_for_pdf(_escape_html(bloco))
            story.append(Paragraph(texto, body_style))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def _escape_html(texto: str) -> str:
    """Escapa caracteres especiais para segurança no XML/HTML do PDF."""
    return (
        texto.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    )


def _bold_for_pdf(texto: str) -> str:
    """Converte pares de `**` em `<b>` (reportlab); remove marcadores não pareados."""
    texto = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", texto, flags=re.DOTALL)
    return texto.replace("**", "")
